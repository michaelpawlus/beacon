"""Resume rendering to various output formats for Beacon Phase 3."""

from pathlib import Path

from beacon.materials.resume import TailoredResume


def render_markdown(resume: TailoredResume) -> str:
    """Render a tailored resume as markdown."""
    return resume.markdown


def render_docx(resume: TailoredResume, output_path: str | Path) -> Path:
    """Render a tailored resume as a DOCX file.

    Requires python-docx: pip install beacon[docs]
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install beacon[docs]")

    doc = Document()

    # Parse markdown into docx
    lines = resume.markdown.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)

    path = Path(output_path)
    doc.save(str(path))
    return path


def render_pdf(resume: TailoredResume, output_path: str | Path) -> Path:
    """Render a tailored resume as a PDF file.

    Requires fpdf2: pip install beacon[docs]
    """
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 not installed. Run: pip install beacon[docs]")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    lines = resume.markdown.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
        if line.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, line[2:], new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 8, line[3:], new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("### "):
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(0, 7, line[4:], new_x="LMARGIN", new_y="NEXT")
        elif line.startswith("- ") or line.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(5)
            pdf.multi_cell(0, 5, f"  {line}")
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, line)

    path = Path(output_path)
    pdf.output(str(path))
    return path
